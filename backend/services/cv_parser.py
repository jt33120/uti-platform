import pdfplumber
import io
import re
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract and clean text from a PDF file.
    Uses pdfplumber for accurate text extraction with layout awareness.
    """
    text_parts = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Extract text with layout preservation
            page_text = page.extract_text(
                x_tolerance=3,
                y_tolerance=3,
                layout=True,
                x_density=7.25,
                y_density=13,
            )
            if page_text:
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

            # Also extract tables if any
            tables = page.extract_tables()
            for table in tables:
                if table:
                    table_text = "\n".join(
                        " | ".join(str(cell) if cell else "" for cell in row)
                        for row in table
                        if any(cell for cell in row)
                    )
                    if table_text.strip():
                        text_parts.append(f"[Table]\n{table_text}")

    raw_text = "\n\n".join(text_parts)
    return clean_text(raw_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file (paragraphs + table cells)."""
    from docx import Document  # imported lazily so the dep is only needed when used

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


# Titres de feuilles servant de catalogues de référence / sources de listes
# déroulantes dans les modèles de marché (CCTP/CRT/AF). Leur contenu est du
# bruit pour l'extraction d'un AO : on les conserve mais reléguées en annexe
# (et tronquées) pour ne pas noyer les vraies données de la consultation.
_XLSX_REFERENCE_HINTS = (
    "liste", "déroulante", "deroulante", "catégorie", "categorie",
    "référence", "reference", "annexe", "nomenclature",
)
_XLSX_REF_MAX_CHARS = 1500

# Cellules/lignes de boilerplate des gabarits de marché (instructions de saisie,
# indices de listes déroulantes) : ce ne sont PAS des données de l'AO — on les
# retire pour ne pas polluer l'extraction envoyée au LLM.
_XLSX_CELL_NOISE = {"liste déroulante", "liste deroulante"}
_XLSX_LINE_NOISE = (
    "cellules à renseigner obligatoirement",
    "ce planning sera à confirmer",
    "ce n'est pas obligatoire de le diffuser",
    "il est obligatoire d'indiquer ce montant",
)


# Champs clés des modèles de marché (label → alias possibles, en minuscules).
# Pré-extraits de façon déterministe et placés EN TÊTE pour guider le LLM.
_XLSX_KEY_LABELS = [
    ("Référence", ("références de la consultation", "référence de la consultation")),
    ("Objet", ("objet de la consultation", "objet de la prestation", "objet du marché")),
    ("Direction", ("direction/support technique", "direction / support technique")),
    ("Interlocuteur", ("interlocuteur technique",)),
    ("Catégorie", ("catégorie concernée",)),
    ("UO", ("uo concernée",)),
    ("Profil", ("profil",)),
    ("Lieu", ("lieu de la prestation",)),
    ("Durée", ("durée du marché",)),
    ("Montant (total)", ("montant maximum pour la durée", "valeur estimée du marché")),
    ("Date limite de remise des offres", ("date de limite de remise des offres", "date limite de remise des offres")),
]


def _xlsx_key_fields(lines: list[str]) -> str:
    """Extrait les champs clés d'un modèle de marché (lignes 'Label : | valeur')
    et les renvoie sous forme d'un bloc synthétique. '' si rien trouvé."""
    found: dict[str, str] = {}
    for line in lines:
        parts = [p.strip() for p in line.split("|")]
        if len(parts) < 2:
            continue
        label = parts[0].rstrip(" :").lower()
        value = " ".join(p for p in parts[1:] if p).strip()
        if not value:
            continue
        for name, aliases in _XLSX_KEY_LABELS:
            if name in found:
                continue
            if any(label == a or label.startswith(a) for a in aliases):
                found[name] = value[:400]
                break
    if not found:
        return ""
    ordered = [f"{name} : {found[name]}" for name, _ in _XLSX_KEY_LABELS if name in found]
    return "[Champs clés extraits du modèle de marché]\n" + "\n".join(ordered)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """
    Extract text from a .xlsx workbook: every sheet, each non-empty row rendered
    as ' | '-joined cells. Useful for AO specs delivered as a spreadsheet
    (cahier des charges Excel). Lazy import so openpyxl is only needed on use.

    Les feuilles « catalogue / listes déroulantes » (ex. liste des UO, des sites,
    des catégories) sont reléguées en annexe et tronquées : sur un modèle de
    marché type AGIRC-ARRCO elles représentent l'essentiel du volume mais aucune
    donnée de l'AO, et placées en tête elles noyaient les vrais champs (objet,
    références, dates, lieu, budget) que l'IA doit extraire.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    main_parts: list[str] = []
    main_lines: list[str] = []
    ref_parts: list[str] = []
    try:
        for ws in wb.worksheets:
            title = (ws.title or "").lower()
            is_ref = any(h in title for h in _XLSX_REFERENCE_HINTS)
            lines: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(c).strip() for c in row
                    if c is not None and str(c).strip()
                    and str(c).strip().lower() not in _XLSX_CELL_NOISE
                ]
                if not cells:
                    continue
                line = " | ".join(cells)
                if any(n in line.lower() for n in _XLSX_LINE_NOISE):
                    continue
                lines.append(line)
            if not lines:
                continue
            block = f"[Feuille : {ws.title}]\n" + "\n".join(lines)
            if is_ref:
                ref_parts.append(block[:_XLSX_REF_MAX_CHARS])
            else:
                main_parts.append(block)
                main_lines.extend(lines)
    finally:
        try:
            wb.close()
        except Exception:
            pass

    parts: list[str] = []
    key_block = _xlsx_key_fields(main_lines)
    if key_block:
        parts.append(key_block)
    parts.extend(main_parts)
    if ref_parts:
        parts.append(
            "[Listes de référence / valeurs possibles — annexes, à ignorer pour l'extraction des champs de l'AO]"
        )
        parts.extend(ref_parts)
    return clean_text("\n".join(parts))


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace while preserving structure
    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        line = line.strip()
        # Skip lines with only special chars or very short noise
        if len(line) > 1 or line.isalnum():
            cleaned_lines.append(line)

    # Collapse multiple blank lines into one
    result = "\n".join(cleaned_lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    result = re.sub(r" {2,}", " ", result)

    return result.strip()


def extract_cv_metadata(text: str) -> dict:
    """
    Try to extract basic metadata from CV text for display purposes.
    Not used for scoring — just for UI enrichment.
    """
    metadata = {
        "email": None,
        "phone": None,
        "estimated_pages": text.count("[Page"),
    }

    # Email
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
    emails = re.findall(email_pattern, text)
    if emails:
        metadata["email"] = emails[0]

    # Phone (French format)
    phone_pattern = r"(?:(?:\+|00)33|0)\s*[1-9](?:[\s.-]*\d{2}){4}"
    phones = re.findall(phone_pattern, text)
    if phones:
        metadata["phone"] = phones[0]

    return metadata
